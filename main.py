from fastapi import FastAPI, File, UploadFile

from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO
from functions.shingling_algorithm import ngrams, preprocess_text
from pathlib import Path

from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from databases import Database
from database.models import Base
from database.models import Shingles

from database.find_similar_document_db import find_similar_documents
from database.config import host, user, password, db_name
from database.insert_shingle_db import insert_shingles


# URL для подключения к базе данных PostgreSQL
DATABASE_URL = "postgresql://postgres:1234@127.0.0.1/PlagiarismDetector"

# асинхронное подключение к базе данных через databases
database = Database(DATABASE_URL)

# синхронный движок для SQLAlchemy
engine = create_engine(DATABASE_URL)
SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)


app = FastAPI()

# Открытие соединения с базой данных при старте приложения
@app.on_event("startup")
async def startup():
    await database.connect()


# Закрытие соединения с базой данных при остановке приложения
@app.on_event("shutdown")
async def shutdown():
    await database.disconnect()



# принимает файл(docx, pdf, txt) и выводит информацию о файле и процент плагиата
@app.post("/uploadfile/")
async def plagiarism_percentage(file: UploadFile):
    contents = await file.read()
    path = Path(file.filename)

    # document_text = file_typer(contents, path)

    if path.suffix == ".docx":
        document = Document(BytesIO(contents))
        document_text = "\n".join([para.text for para in document.paragraphs])


    elif path.suffix == ".pdf":
        pdf_reader = PdfReader(BytesIO(contents))
        document_text = ""

        for page in pdf_reader.pages:
            document_text += page.extract_text()
    

    elif path.suffix == ".txt":
        try:
            document_text = contents.decode("utf-8")
        except UnicodeDecodeError:
            document_text = contents.decode("windows-1251")

    shingles_list = ngrams(preprocess_text(document_text), 5)

    similars = find_similar_documents(shingles_list)
    
    sum_similars = 0
    for similar in similars:
        sum_similars += similar[1]

    plagiarism_percentage = (sum_similars/len(shingles_list))*100

    # return {"filename": file.filename, 'content_type': file.content_type, 'content': shingles_list}
    return {"filename": file.filename, 'content_type': file.content_type, 'similar_docs': similars,
     'len_shingles_list': len(shingles_list), 'plagiarism_percentage': round(plagiarism_percentage, 2)}


# выводит базу данных с шинглами
@app.get("/database/")
async def get_shingles():
    query = "SELECT * FROM shingles"
    shingles = await database.fetch_all(query=query)
    return shingles


@app.post("/database/")
async def add_shingles(file: UploadFile):
    contents = await file.read()
    path = Path(file.filename)

    if path.suffix == ".docx":
        document = Document(BytesIO(contents))
        document_text = "\n".join([para.text for para in document.paragraphs])


    elif path.suffix == ".pdf":
        pdf_reader = PdfReader(BytesIO(contents))
        document_text = ""

        for page in pdf_reader.pages:
            document_text += page.extract_text()
    

    elif path.suffix == ".txt":
        try:
            document_text = contents.decode("utf-8")
        except UnicodeDecodeError:
            document_text = contents.decode("windows-1251")

    shingles_list = ngrams(preprocess_text(document_text), 5)

    name = Path(file.filename).stem

    insert_shingles(name, shingles_list)